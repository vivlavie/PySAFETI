from openpyxl import load_workbook
import math
import numpy as np
import matplotlib.pyplot as plt
from scipy.interpolate import interp1d
import scipy.ndimage
from scipy.ndimage.filters import gaussian_filter
import dill

import matplotlib.pylab as pltparam
from matplotlib.patches import Ellipse

pltparam.rcParams["figure.figsize"] = (8,3)
pltparam.rcParams['lines.linewidth'] = 2
# pltparam.rcParams['lines.color'] = 'r'
pltparam.rcParams['axes.grid'] = True 


import re
def slugify(value):
    """
    Normalizes string, converts to lowercase, removes non-alpha characters,
    and converts spaces to hyphens.
    """
    # import unicodedata
    # value = unicodedata.normalize('NFKD', value).encode('ascii', 'ignore')
    # value = unicode(re.sub('[^\w\s-]', '', value).strip().lower())
    # value = unicode(re.sub('[-\s]+', '-', value))
    value = re.sub('[^\w\s-]', '', value).strip().lower()
    value = re.sub('[-\s]+', '-', value)
    return value

# element_dump_filename = 'Bv06_offloading_dump'
# folder = 'Offloading'
# XDir = -1 #From Stern to Bow

element_dump_filename = 'Bv06_dump'
folder = 'ProcessArea'
XDir = 1 #From Bow to Stern

# element_dump_filename = 'Bv06_utility_dump'
# folder = 'Utility'
# XDir = 1 #From Bow to Stern

# element_dump_filename = 'Bv06_hull_dump'
# folder = 'Hull'
# XDir = 1 #From Bow to Stern


with open(element_dump_filename,'rb') as element_dump:
    lEvent = dill.load(element_dump)


for e in lEvent:    
# for e in [lEvent[20]]:    
    # if ("-L" in e.Key) and ("LM_EO" in e.Hole) and (e.Weather == "7.7D"):    
        Key = e.Key
        X = e.X
        Y = e.Y
        
        fig,ax = plt.subplots()
        if e.Deck == 'A':
            img = plt.imread("RubyFPSODeckA.jpg")
        elif (e.Deck == 'B') or (e.Deck == 'C'):
            img = plt.imread("RubyFPSODeckB.jpg")
        elif e.Deck == 'Hull':
            # print('Hull deck?')
            img = plt.imread("RubyFPSOHullDeck.jpg")
        ax.imshow(img, extent=[-11, 252, -32.43, 50.19])

        #Plot Dispersion - LFL Fraction
        # if e.Dispersion != None:
        #     DfMin = e.Dispersion.DfMin
        #     DfMax = e.Dispersion.DfMax
        #     Wf = e.Dispersion.Wf
        #     if Wf == ' ':
        #         Wf = e.Dispersion.FFWidth
        #     AxisLongf = (DfMax - DfMin)
        #     AxisShortf = Wf
        #     DXf = (DfMax + DfMin)*0.5
        #     cloudf = Ellipse((X-DXf,Y),AxisLongf,AxisShortf,0)
        #     ax.add_artist(cloudf)
        #     cloudf.set_alpha(0.5)
        #     cloudf.set_facecolor('green')

        #Plot Dispersion - LFL
        if (e.Dispersion != None) and (e.Dispersion.DMin != ' '):
            DMin = e.Dispersion.DMin
            DMax = e.Dispersion.DMax
            W = e.Dispersion.W    
            AxisLong = (DMax - DMin)
            AxisShort = W
            DX = (DMax + DMin)*0.5            
            cloud = Ellipse((X-XDir*DX,Y),AxisLong,AxisShort,0)
            # ax.add_artist(cloud)
            ax.add_patch(cloud)
            cloud.set_alpha(0.9)
            cloud.set_facecolor('Magenta')
            cloud.set(label='LFL')
        
        #Plot Dispersion - Flash Fire
        if (e.Dispersion != None) and (e.Dispersion.FFMaxDistance != ' '):        
            FFMax = e.Dispersion.FFMaxDistance
            FFW = e.Dispersion.FFWidth
            AxisLong = FFMax
            AxisShort = FFW
            
            cloud = Ellipse((X-XDir*AxisLong*0.5,Y),AxisLong,AxisShort,0)
            # ax.add_artist(cloud)
            ax.add_patch(cloud)
            cloud.set_alpha(0.3)
            cloud.set_facecolor('Yellow')
            cloud.set(label='Flash (50\% LFL)')


        #Plot Jet Fire
        if e.JetFire != None:
            jfl = e.JetFire.Length
            jfw = 0.12*jfl
            JF = np.array([[X,Y],[X-XDir*jfl,Y+0.5*jfw],[X-XDir*jfl,Y-0.5*jfw]])
            jfpatch = plt.Polygon(JF,color='red')
            ax.add_patch(jfpatch)
            jfpatch.set(label='Jet')


        #Plot Early Pool Fire
        if e.EarlyPoolFire != None:
            EPFD = e.EarlyPoolFire.Diameter
            EPF = Ellipse((X,Y),EPFD,EPFD,0)
            # ax.add_artist(EPF)
            ax.add_patch(EPF)
            EPF.set_alpha(0.5)
            EPF.set_facecolor('blue')
            EPF.set(label='Early Pool')

        if e.LatePoolFire != None:
            LPFD = e.LatePoolFire.Diameter
            LPF = Ellipse((X,Y),LPFD,LPFD,0)
            # ax.add_artist(LPF)
            ax.add_patch(LPF)
            LPF.set_alpha(0.5)
            LPF.set_facecolor('cyan')
            LPF.set(label='Late Pool')
        # ax.clabel(cs2,fmt='%.1e',colors='k',fontsize=14)
        # ax.set_aspect('equal')
        # ax.xaxis.set_major_locator(plt.FixedLocator([120, 141, 168, 193]))
        # ax.xaxis.set_major_formatter(plt.FixedFormatter(['2/3','3/4','4/5','S05']))
        # ax.yaxis.set_major_locator(plt.FixedLocator([-27, -3.1, 3.1, 27]))
        # ax.yaxis.set_major_formatter(plt.FixedFormatter(['ER_S','Tray_S','Tray_P','ER_P']))

        ax.legend()
        plt.title(Key)
        # fn = 'C\\ProcessArea\\C_'+slugify(Key)+'.png'
        fn = 'C\\'+folder+'\\C_'+slugify(Key)+'.png'

        fig.savefig(fn)
        # plt.show()
        plt.close()



from docx import Document
from docx.shared import Cm, RGBColor, Pt
from docx.text.run import Font, Run
# from docx.dml.color import ColorFormat

document = Document()


""" runa = document.add_paragraph('Appendix').add_run()
fonta = runa.font
fonta.name = 'Frutiger LT 55'
fonta.size = Pt(18)
fonta.color.rgb = RGBColor(59, 142, 222)


runb = document.add_paragraph('PHAST Analysis - ' + folder).add_run()
fontb = runb.font
fontb.size = Pt(30)
fontb.color.rgb = RGBColor(31, 73, 125)

document.add_page_break()

fontn = document.styles['Normal'].font
fontn.size = Pt(9)
fontn.name = 'Frutiger LT 45'
fontn.color.rgb = RGBColor(31, 73, 125)

document.add_heading(IS, level=1)
document.add_paragraph('Figure ')

document.save('C_'+folder+'.docx') """


""" p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
 """

fontn = document.styles['Normal'].font
fontn.size = Pt(9)
fontn.name = 'Frutiger LT 45'
fontn.color.rgb = RGBColor(31, 73, 125)

document.add_heading('PHAST Analysis - ' + folder, level=0)

ISs = []
for e in lEvent:
    pvis,hole,weather = e.Key.split("\\")
    if not (pvis in ISs):
        ISs.append(pvis)
nbr=1
for IS in ISs:
    document.add_heading(IS, level=1)
    for e in lEvent:
        pvis,hole,weather = e.Key.split("\\")
        if IS == pvis:                        
            fn = 'C\\'+folder+'\\C_'+slugify(e.Key)+'.png'
            document.add_picture(fn, width=Cm(15))            
            p = document.add_paragraph('Figure '+str(nbr) + " - " + pvis + " Hole: " + hole + " Weather: " + weather)
            nbr = nbr +1
    document.add_page_break()
document.save('C_'+folder+'.docx')
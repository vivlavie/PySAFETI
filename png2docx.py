from docx import Document
from docx.shared import Inches
from docx.text.run import Font, Run
# from docx.dml.color import ColorFormat

document = Document()
document.add_heading('RUBY FPSO - Time Varying Discharges', 0)

font = document.styles['Normal'].font
font.name = 'Frutiger LT 55'
# font.color = 'Dark Blue, Text 2'
# font.color.rgb = RGBColor(31, 73, 125)

""" p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
 """
nbr=1
for Hole in ['06']:
    document.add_heading('Hole '+ Hole +' mm', level=1)
    # for Hole in ['13':
    for Rel in ['F', 'V','A']:
    # for Rel in ['V','F']:
        document.add_heading('Release Direction: '+Rel, level=2)
        for WD in ['H','S','P','T']:
            document.add_heading('Wind Direction: '+ WD, level=3)
        # for WD in ['H']:
            for WV in ['5','1']:
                scn = Hole + Rel + WD + WV
                fn50 = './Plots/50LFL/'+scn+'_iso_50LFL.jpg'
                fn100 = './Plots/100LFL/'+scn+'_iso_LFL.jpg'
                fncntr = './Plots/Contour/'+scn+'_contour_z.jpg'
                document.add_page_break()
                document.add_picture(fn100, width=Inches(5.6))
                document.add_picture(fn50, width=Inches(5.6))
                document.add_picture(fncntr, width=Inches(5.6))
                p = document.add_paragraph('Figure A.'+str(48+nbr)+' LFL gas cloud of ' + scn +'(Upper:100%%, Middle:50%%, Lower:Contour)')

                nbr = nbr +1
                
# document.add_page_break()

document.save('VolPlotsv02.docx')

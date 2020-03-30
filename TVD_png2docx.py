from openpyxl import load_workbook
import dill
from docx import Document
from docx.shared import Cm, RGBColor, Pt
from docx.text.run import Font, Run
# from docx.dml.color import ColorFormat
import re


def slugify(value):
    """
    Normalizes string, converts to lowercase, removes non-alpha characters,
    and converts spaces to hyphens.
    """
    # import unicodedata
    # value = unicodedata.normalize('NFKD', value).encode('ascii', 'ignore')
    # value = unicode(re.sub('[^\w\s-]', '', value).strip().lower())
    # value = unicode(re.sub('[-\s]+', '-', value))
    value = re.sub('[^\w\s-]', '', value).strip().lower()
    value = re.sub('[-\s]+', '-', value)
    return value

# element_dump_filename = 'Bv06_offloading_dump'
# folder = 'Offloading'
# element_dump_filename = 'Bv06_utility_dump'
# folder = 'Utility'
# element_dump_filename = 'Bv06_hull_dump'
# folder = 'Hull'
element_dump_filename = 'Bv06_dump'
folder = 'ProcessArea'

with open(element_dump_filename,'rb') as element_dump:
    lEvent = dill.load(element_dump)


document = Document()


fontn = document.styles['Normal'].font
fontn.size = Pt(9)
fontn.name = 'Frutiger LT 45'
fontn.color.rgb = RGBColor(31, 73, 125)

document.add_heading('PHAST Discharge Analysis - ' + folder, level=0)

ISs = []
for e in lEvent:
    pvis,hole,weather = e.Key.split("\\")
    if not (pvis in ISs):
        ISs.append(pvis)

nbr=1
for IS in ISs:
    document.add_heading(IS, level=1)
    for e in lEvent:
        pvis,hole,weather = e.Key.split("\\")
        if IS == pvis:                        
            # fn = 'C\\'+folder+'\\C_'+slugify(e.Key)+'.png'
            pngfilename = "TVD-"+pvis+"_"+hole
            fn = slugify(pngfilename)
            document.add_picture(".\\tvd_rev.B\\"+fn+".png", width=Cm(13))            
            p = document.add_paragraph('Figure '+str(nbr) + " - " + pvis + " Hole: " + hole)
            nbr = nbr +1
    document.add_page_break()
document.save('tvd_'+folder+'.docx')
